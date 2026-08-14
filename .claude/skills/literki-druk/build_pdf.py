#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Buduje wersję do druku (PDF/DOCX) bajek o literkach.

Cykl „Bajki o literkach": JEDEN plik = JEDNA część (opowiadania/czesc-NN - <podtytuł>.md).

Domyślnie (bez argumentów): PDF **najnowszej** (najwyższej numerem) części.

Parametry:
  --typ  pdf | docx | oba          (domyślnie: pdf)
  --czesc N | calosc | wszystko    (domyślnie: najnowsza część)
       N        -> jedna część (np. --czesc 5)
       calosc   -> jeden dokument z całością (wszystkie części po kolei)
       wszystko -> każda część osobno + całość

WAŻNE: PDF powstaje CZYSTO w Pythonie (reportlab) — **bez MS Word i bez żadnej drukarki**.
Cechy PDF:
  - klikalny **spis treści** (skacze do części) + numery stron
  - **struktura po lewej** (zakładki/outline): Część -> Scenka (nagłówek ###), klikalna
  - **hasła dla dziecka** (linie z 👉) renderowane jako WYRÓŻNIONE, wyśrodkowane ramki
    (emoji 👉 jest znacznikiem formatu w źródle; w druku zastępuje je wyróżnienie)
  - tekst 12 (justowany), część 20, scenka 14; każda część od nowej strony; numeracja stron
  - czcionka Montserrat, jeśli TTF jest w systemie/`fonts/`; w innym razie Arial (polskie znaki OK)
DOCX (opcjonalnie): python-docx, nagłówki jako style Word + pole spisu treści (Word: Ctrl+A, F9).
"""
import argparse
import re
import sys
import os
from pathlib import Path

FONT = "Montserrat"
SIZE_BODY = 12
SIZE_HASLO = 15
SIZE_CHAPTER = 14
SIZE_PART = 20
SIZE_SERIES = 14

SERIES_TITLE = "Bajki o literkach - czytamy razem"
BOOK_TITLE = "Literkowa Kraina. Wszystkie przygody z literami"
BOOK_BLURB = ("Dwadzieścia interaktywnych bajek do czytania razem: rodzic czyta historię, "
              "a dziecko - magiczne hasła, bez których akcja nie ruszy dalej.")

REPO = Path(__file__).resolve().parents[3]
STORIES_DIR = REPO / "opowiadania"
OUT_DIR = REPO / "druk"
OPISY = STORIES_DIR / "opisy.md"

# pliki w opowiadania/, które NIE są opowiadaniami
META_FILES = {"rejestr.md", "opisy.md", "wzorce.md", "plan.md", "readme.md"}
INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")


# ----------------------------- wspólne -----------------------------

def load_opisy():
    opisy = {}
    if not OPISY.exists():
        return opisy
    cur, buf = None, []
    for line in OPISY.read_text(encoding="utf-8").splitlines():
        m = re.match(r"^##\s+Część\s+(\d+)", line)
        if m:
            if cur is not None:
                opisy[cur] = " ".join(buf).strip()
            cur = int(m.group(1)); buf = []
        elif cur is not None and line.strip() and not line.startswith("#"):
            buf.append(line.strip())
    if cur is not None:
        opisy[cur] = " ".join(buf).strip()
    return opisy


def part_files():
    """Zwraca {numer_części: ścieżka} dla plików opowiadania/czesc-*.md (pomija meta)."""
    out = {}
    for p in STORIES_DIR.glob("czesc-*.md"):
        if p.name.lower() in META_FILES:
            continue
        m = re.match(r"czesc-(\d+)", p.name)
        if m:
            out[int(m.group(1))] = p
    return out


def part_file(c):
    return part_files().get(c)


def list_parts():
    return sorted(part_files().keys())


def subtitle_of(md_path):
    first = md_path.read_text(encoding="utf-8").splitlines()[0]
    m = re.match(r"^#\s*(?:Część\s*)?\d+\s*[-—]\s*(.+)$", first)
    if m:
        return m.group(1).strip()
    # awaryjnie z nazwy pliku "czesc-NN - <podtytuł>.md"
    name = md_path.name.rsplit(".md", 1)[0]
    return name.split(" - ", 1)[1] if " - " in name else name


def part_title(c, md_path):
    return f"Część {c} - {subtitle_of(md_path)}"


def story_lines(md_path):
    """Zwraca listę ('title'|'chapter'|'haslo'|'quote'|'body', tekst)."""
    out = []
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            out.append(("title", line[2:].strip()))
        elif line.startswith("### "):
            out.append(("chapter", line[4:].strip()))
        elif line.startswith("## "):
            out.append(("chapter", line[3:].strip()))
        elif line.lstrip().startswith("👉"):
            out.append(("haslo", line.lstrip()[1:].strip()))
        elif line.startswith("> "):
            body = line[2:].strip()
            if body.startswith("👉"):
                out.append(("haslo", body[1:].strip()))
            else:
                out.append(("quote", body))
        else:
            out.append(("body", line))
    return out


# ============================ PDF (reportlab) ============================

def _find_font_family():
    dirs = [REPO / "fonts", Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts"]
    def find(names):
        for d in dirs:
            for n in names:
                p = d / n
                if p.exists():
                    return str(p)
        return None
    mont = (find(["Montserrat-Regular.ttf", "Montserrat.ttf"]),
            find(["Montserrat-Bold.ttf", "Montserrat-SemiBold.ttf"]),
            find(["Montserrat-Italic.ttf"]),
            find(["Montserrat-BoldItalic.ttf"]))
    if mont[0] and mont[1]:
        return mont[0], mont[1], mont[2] or mont[0], mont[3] or mont[1], "Montserrat"
    ar = (find(["arial.ttf", "Arial.ttf"]), find(["arialbd.ttf"]),
          find(["ariali.ttf"]), find(["arialbi.ttf"]))
    if ar[0] and ar[1]:
        return ar[0], ar[1], ar[2] or ar[0], ar[3] or ar[1], "Arial"
    return None


def _register_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    fam = _find_font_family()
    if not fam:
        print("UWAGA: brak Montserrat/Arial TTF — PDF użyje Helvetica (polskie znaki mogą być gorsze).",
              file=sys.stderr)
        return "Helvetica", "Helvetica-Bold", "Helvetica-Oblique", "Helvetica-BoldOblique", "Helvetica"
    reg, bold, ital, bi, label = fam
    base = "LK"
    pdfmetrics.registerFont(TTFont(base, reg))
    pdfmetrics.registerFont(TTFont(base + "-B", bold))
    pdfmetrics.registerFont(TTFont(base + "-I", ital))
    pdfmetrics.registerFont(TTFont(base + "-BI", bi))
    pdfmetrics.registerFontFamily(base, normal=base, bold=base + "-B",
                                  italic=base + "-I", boldItalic=base + "-BI")
    if label != "Montserrat":
        print(f"Info: Montserrat nie znaleziony — PDF użyje czcionki {label}.", file=sys.stderr)
    return base, base + "-B", base + "-I", base + "-BI", label


def _rml(text):
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def build_pdf(targets, opisy):
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                    Spacer, PageBreak)
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT

    REG, BOLD, ITAL, BI, label = _register_fonts()

    body = ParagraphStyle("body", fontName=REG, fontSize=SIZE_BODY, leading=SIZE_BODY * 1.4,
                          alignment=TA_JUSTIFY, spaceAfter=6)
    quote = ParagraphStyle("quote", parent=body, leftIndent=0.6 * cm, rightIndent=0.6 * cm)
    haslo = ParagraphStyle("haslo", fontName=BOLD, fontSize=SIZE_HASLO, alignment=TA_CENTER,
                           leading=SIZE_HASLO * 1.35, spaceBefore=8, spaceAfter=10,
                           textColor=colors.HexColor("#1b3a6b"),
                           backColor=colors.HexColor("#fff3c4"),
                           borderColor=colors.HexColor("#e0a800"), borderWidth=1.2,
                           borderPadding=(7, 7, 7), borderRadius=6)
    czesc = ParagraphStyle("Czesc", fontName=BOLD, fontSize=SIZE_PART, alignment=TA_CENTER,
                           spaceBefore=0, spaceAfter=16, leading=SIZE_PART * 1.2)
    rozdz = ParagraphStyle("Rozdzial", fontName=BOLD, fontSize=SIZE_CHAPTER, alignment=TA_LEFT,
                           spaceBefore=12, spaceAfter=6, leading=SIZE_CHAPTER * 1.2)
    series = ParagraphStyle("series", fontName=REG, fontSize=SIZE_SERIES, alignment=TA_CENTER, spaceAfter=6)
    bigt = ParagraphStyle("bigt", fontName=BOLD, fontSize=28, alignment=TA_CENTER, spaceAfter=18,
                          leading=28 * 1.2)
    blurb = ParagraphStyle("blurb", fontName=ITAL, fontSize=SIZE_BODY, alignment=TA_CENTER,
                           leading=SIZE_BODY * 1.4, leftIndent=1.2 * cm, rightIndent=1.2 * cm)
    tochdr = ParagraphStyle("tochdr", fontName=BOLD, fontSize=SIZE_CHAPTER, spaceAfter=10)
    toc0 = ParagraphStyle("toc0", fontName=BOLD, fontSize=12, leading=18, spaceBefore=6)

    class LiterkiDoc(BaseDocTemplate):
        def __init__(self, path):
            super().__init__(path, pagesize=A4, title=BOOK_TITLE)
            self._h = 0
            fr = Frame(2.5 * cm, 2.2 * cm, A4[0] - 4.5 * cm, A4[1] - 4.4 * cm, id="b")
            self.addPageTemplates([PageTemplate(id="main", frames=[fr], onPage=self._foot)])

        def build(self, flowables, **kw):
            self._h = 0
            return super().build(flowables, **kw)

        def _foot(self, canvas, doc):
            canvas.saveState(); canvas.setFont(REG, 9)
            canvas.drawCentredString(A4[0] / 2, 1.2 * cm, str(doc.page))
            canvas.restoreState()

        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                sn = flowable.style.name
                if sn in ("Czesc", "Rozdzial"):
                    text = flowable.getPlainText()
                    level = {"Czesc": 0, "Rozdzial": 1}[sn]
                    key = "h%d" % self._h; self._h += 1
                    self.canv.bookmarkPage(key)
                    self.canv.addOutlineEntry(text, key, level=level, closed=(level > 0))
                    if sn == "Czesc":
                        self.notify("TOCEntry", (0, text, self.page, key))

    def story_flow(md, c):
        fl = [PageBreak()]
        for kind, text in story_lines(md):
            if kind == "title":
                fl.append(Paragraph(_rml(part_title(c, md)), czesc))
            elif kind == "chapter":
                fl.append(Paragraph(_rml(text), rozdz))
            elif kind == "haslo":
                fl.append(Paragraph(_rml(text), haslo))
            elif kind == "quote":
                fl.append(Paragraph(_rml(text), quote))
            else:
                fl.append(Paragraph(_rml(text), body))
        return fl

    def make_toc():
        t = TableOfContents()
        t.levelStyles = [toc0]
        return t

    OUT_DIR.mkdir(exist_ok=True)
    made = []
    for kind, arg in targets:
        if kind == "single":
            c = arg; md = part_file(c)
            if not md:
                print(f"UWAGA: brak pliku części {c:02d}", file=sys.stderr); continue
            out = OUT_DIR / f"Literki - Czesc {c:02d}.pdf"
            flow = [Spacer(1, 5 * cm), Paragraph(SERIES_TITLE, series),
                    Paragraph(part_title(c, md), bigt)]
            if opisy.get(c):
                flow.append(Paragraph(_rml(opisy[c]), blurb))
            flow += story_flow(md, c)
            LiterkiDoc(str(out)).multiBuild(flow)
            print(f"Zapisano PDF: {out}"); made.append(out)
        else:
            cols = arg; out = OUT_DIR / "Literki - calosc.pdf"
            flow = [Spacer(1, 5 * cm), Paragraph(SERIES_TITLE, series),
                    Paragraph(BOOK_TITLE, bigt), Paragraph(BOOK_BLURB, blurb),
                    PageBreak(), Paragraph("Spis treści", tochdr), make_toc(), PageBreak()]
            for c in cols:
                md = part_file(c)
                if not md:
                    continue
                flow += story_flow(md, c)
            LiterkiDoc(str(out)).multiBuild(flow)
            print(f"Zapisano PDF: {out}"); made.append(out)
    return made


# ============================ DOCX (python-docx) ============================

def build_docx(targets, opisy):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _rfonts(rpr):
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), FONT)

    def add_run(p, text, size, bold=False, italic=False, color=None):
        r = p.add_run(text); r.font.name = FONT; _rfonts(r._element.get_or_add_rPr())
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color:
            r.font.color.rgb = color
        lang = OxmlElement("w:lang"); lang.set(qn("w:val"), "pl-PL")
        r._element.get_or_add_rPr().append(lang)
        return r

    def add_inline(p, text, size, bold=False, italic=False):
        pos = 0
        for m in INLINE.finditer(text):
            if m.start() > pos:
                add_run(p, text[pos:m.start()], size, bold, italic)
            if m.group(1) is not None:
                add_run(p, m.group(1), size, True, italic)
            else:
                add_run(p, m.group(2), size, bold, True)
            pos = m.end()
        if pos < len(text):
            add_run(p, text[pos:], size, bold, italic)

    def setup():
        doc = Document()
        nrm = doc.styles["Normal"]; nrm.font.name = FONT; nrm.font.size = Pt(SIZE_BODY)
        nrm.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        nrm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        nrm.paragraph_format.line_spacing = 1.3
        nrm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        for name, size in (("Heading 1", SIZE_PART), ("Heading 2", SIZE_CHAPTER)):
            st = doc.styles[name]; st.font.name = FONT; st.font.size = Pt(size)
            st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
            _rfonts(st.element.get_or_add_rPr()); st.paragraph_format.keep_with_next = True
        s = doc.sections[0]
        s.page_height = Cm(29.7); s.page_width = Cm(21.0)
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
        upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true"); doc.settings.element.append(upd)
        fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = fp.add_run()
        for typ, txt in (("begin", None), ("instr", "PAGE"), ("end", None)):
            if typ == "instr":
                el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = txt
            else:
                el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), typ)
            rr._r.append(el)
        rr.font.name = FONT; rr.font.size = Pt(10)
        return doc

    def _shade(p, hexcolor):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
        pPr.append(shd)

    def hding(doc, text, level, size, center=False, pb=False):
        p = doc.add_paragraph(style=f"Heading {level}")
        if center:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if pb:
            p.paragraph_format.page_break_before = True
        add_inline(p, text, size, bold=True)

    def bodyp(doc, text, italic=False, indent=None):
        p = doc.add_paragraph(); add_inline(p, text, SIZE_BODY, italic=italic)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)

    def haslop(doc, text):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
        _shade(p, "FFF3C4")
        clean = text.replace("**", "")
        add_run(p, clean, SIZE_HASLO, bold=True, color=RGBColor(0x1b, 0x3a, 0x6b))

    def toc_field(doc):
        p = doc.add_paragraph(); run = p.add_run()
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = 'TOC \\o "1-1" \\h \\z \\u'
        sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
        tr = OxmlElement("w:r"); tt = OxmlElement("w:t"); tt.text = "Spis treści — w Wordzie: Ctrl+A, potem F9."; tr.append(tt)
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        run._r.append(b); run._r.append(ins); run._r.append(sep); p._p.append(tr)
        p.add_run()._r.append(e)

    def title_page(doc, big, bl=None):
        for _ in range(5):
            doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, SERIES_TITLE, SIZE_SERIES)
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, big, 28, bold=True)
        if bl:
            pb = doc.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pb.paragraph_format.left_indent = Cm(1.5); pb.paragraph_format.right_indent = Cm(1.5)
            add_run(pb, bl, SIZE_BODY, italic=True)

    def render(doc, md, c, part_pb=True):
        for kind, text in story_lines(md):
            if kind == "title":
                hding(doc, part_title(c, md), 1, SIZE_PART, center=True, pb=part_pb)
            elif kind == "chapter":
                hding(doc, text, 2, SIZE_CHAPTER)
            elif kind == "haslo":
                haslop(doc, text)
            elif kind == "quote":
                bodyp(doc, text, italic=True, indent=0.8)
            else:
                bodyp(doc, text)

    OUT_DIR.mkdir(exist_ok=True); made = []
    for kind, arg in targets:
        doc = setup()
        if kind == "single":
            c = arg; md = part_file(c)
            if not md:
                print(f"UWAGA: brak pliku części {c:02d}", file=sys.stderr); continue
            out = OUT_DIR / f"Literki - Czesc {c:02d}.docx"
            title_page(doc, part_title(c, md), opisy.get(c)); toc_field(doc)
            render(doc, md, c, part_pb=True)
        else:
            cols = arg; out = OUT_DIR / "Literki - calosc.docx"
            title_page(doc, BOOK_TITLE, BOOK_BLURB); toc_field(doc)
            for c in cols:
                md = part_file(c)
                if not md:
                    continue
                render(doc, md, c, part_pb=True)
        doc.save(str(out)); print(f"Zapisano DOCX: {out}"); made.append(out)
    return made


# ============================ CLI ============================

def resolve_targets(czesc, cols):
    if czesc is None:
        return [("single", max(cols))]
    cz = str(czesc).strip().lower()
    if cz == "wszystko":
        return [("single", c) for c in cols] + [("all", cols)]
    if cz in ("calosc", "całość", "all"):
        return [("all", cols)]
    return [("single", int(cz))]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--typ", choices=["pdf", "docx", "oba"], default="pdf")
    ap.add_argument("--czesc", default=None, help="N | calosc | wszystko (domyślnie: najnowsza część)")
    ap.add_argument("--out", default=None)
    args = ap.parse_args()
    cols = list_parts()
    if not cols:
        print("Brak części w opowiadania/ (pliki czesc-*.md)", file=sys.stderr); return
    opisy = load_opisy()
    targets = resolve_targets(args.czesc, cols)
    if args.typ in ("docx", "oba"):
        build_docx(targets, opisy)
    if args.typ in ("pdf", "oba"):
        build_pdf(targets, opisy)


if __name__ == "__main__":
    main()
